#!/usr/bin/env python3
# LMS 외부 서비스 결정 — 심플 워드 문서(.docx)
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
FONT = "맑은 고딕"

doc = Document()

# 기본 폰트(한글 포함)
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
st.paragraph_format.space_after = Pt(4)
st.paragraph_format.line_spacing = 1.25

GRAY = RGBColor(0x66, 0x66, 0x66)

def _font(run, size=None, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size: run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = color

def title(text, sub=None):
    p = doc.add_paragraph()
    r = p.add_run(text); _font(r, 19, True)
    p.paragraph_format.space_after = Pt(2)
    if sub:
        ps = doc.add_paragraph(); rs = ps.add_run(sub); _font(rs, 10, False, GRAY)
        ps.paragraph_format.space_after = Pt(10)

def h(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text); _font(r, 13, True)

def body(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix); _font(r, 10.5, True)
    r = p.add_run(text); _font(r, 10.5)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); _font(r, 10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def note(text):  # 회색 작은 참고문구
    p = doc.add_paragraph(); r = p.add_run(text); _font(r, 9.5, False, GRAY)
    return p

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(hd); _font(r, 9.5, True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(val); _font(r, 9.5, i > 0 and row is rows[0])
    return t

# ─────────────────────────────────────────────────────────
title("LMS 외부 서비스 검토", "자체 개발 외 외부 유료 서비스 · 대표·이해관계자 결정용")

h("결정 요약")
bullet("영상 호스팅 — 결정 필요 (1안 DRM 호스팅 vs 2안 YouTube)")
bullet("클리닉 화상 — 구글 미트 (확정 제안)")
bullet("문의·상담·전화 — 채널톡 (확정 제안)")
body("이번에 실제로 결정할 것은 영상 호스팅(1안 vs 2안) 하나입니다.")

# 1. 영상 호스팅
h("1. 영상 호스팅  [결정 필요]")
body("다운로드 불가 / 콘솔로 링크를 찾아도 재생 불가(= DRM 필수) / 워터마크(가점) / 4K 원본 · LMS 임베드", bold_prefix="요건 · ")
body("영상 자체를 암호화. 개발자도구로 주소를 빼내도 재생 불가, 열쇠는 로그인한 그 학생·기기에만 발급.", bold_prefix="DRM · ")
body("사람이 하던 반복 클릭을 봇이 대신. 유튜브 권한 부여/회수를 자동 처리.", bold_prefix="RPA · ")

body("영상 30개 × 1시간 · 각자 전체를 월 2회 시청 · 괄호는 1인당(단위 원/월)", bold_prefix="1안 · DRM 호스팅 (Mux / VdoCipher / 네이버클라우드) — ")
p = doc.add_paragraph(); r = p.add_run("4K"); _font(r, 10, True)
table(["서비스", "50명 (1인당)", "100명 (1인당)"], [
    ["Mux", "51만 (10,200)", "131만 (13,100)"],
    ["VdoCipher", "173만 (34,600)", "330만 (33,000)"],
    ["네이버클라우드", "164만 (32,800)", "281만 (28,100)"],
])
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); r = p.add_run("1080p"); _font(r, 10, True)
table(["서비스", "50명 (1인당)", "100명 (1인당)"], [
    ["Mux", "26만 (5,200)", "50만 (5,000)"],
    ["VdoCipher", "58만 (11,600)", "87만 (8,700)"],
    ["네이버클라우드", "58만 (11,600)", "89만 (8,900)"],
])
doc.add_paragraph()
bullet("기능(DRM·LMS 임베드)은 3사 동일. 워터마크만 VdoCipher 포함.")
bullet("비용은 대부분 '스트리밍'(시청량 × 화질에 비례). 그래서 4K가 1080p의 2~3배.")

body("비용 0원 · 4K 무료 · RPA 검증 완료(부여·회수 실동작). 단 아래 제약이 있음.", bold_prefix="2안 · YouTube (RPA) — ")
bullet("영상당 32계정 한도 → 100명이면 영상 사본 4개로 분할 필요")
bullet("유튜브 코드 변경 시 공유 중단(복구 약 1일) · LMS 임베드 불가 · 워터마크 없음")

# 2. 클리닉
h("2. 클리닉 화상  [확정 제안 · 구글 미트]")
body("신청 → 미트 링크 자동 생성 → 진행(녹화·전사·AI요약 ON) → 구글 드라이브 저장 → LMS 기록(링크 + AI요약)", bold_prefix="흐름 · ")
bullet("출결 자동 판정(입장·퇴장 시각), 전사·AI요약 한국어 지원, 학생은 무료 참여")
bullet("비용: 호스트(조교) 계정당 $14/월 × 동시 진행 조교 수")

# 3. 채널톡
h("3. 문의·상담·전화  [확정 제안 · 채널톡]")
bullet("문의 통합: LMS 위젯 + 카카오톡 + 인스타 DM을 한 곳에서 응대")
bullet("전화: 발신 + 자동 녹음 + AI 요약, 학생별 통화 이력 (개인 폰은 기록이 안 남음)")
bullet("결석생 상담 명단 자동 추출 · 통화 횟수 표시, CRM·헬프센터")
bullet("비용: 월 27,000원 + 전화번호 5,000원/월 + 통화료(무선 72원/분) → 대략 월 5~6만원")

# 참고
h("참고 · 결정 아님 (이미 사용 / 방향 확정)")
bullet("교재 결제 · 결제선생: 이미 사용 중, LMS에 연동/임베드만. 신규 비용 없음.")
bullet("알림톡 · 솔라피: 성적·클리닉·결제 등 자동 발송(서버 직접). 건당 6.5~13원.")

out = os.path.join(HERE, "LMS 외부서비스 결정.docx")
doc.save(out)
print("saved", out)
